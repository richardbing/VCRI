#!/usr/bin/env python
# encoding: utf-8
import os
from docx import Document
from docx.document import Document as _Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

class docxParser(object):
    def __init__(self, path, *args, **kwargs):
        self.path = path
        self.document = Document(path)
        self.tables = self.document.tables
        self.sysRqmtTable = []
        self.srsRqmtTable = []
        self.srsRqmtTable1 = []
        self.systemRequirementList = []
        self.srsRequirementList = []
        self.textStringList = []
        self.getRequirementTable()
        self.getSystemRequirementList()
        self.getSrsRequirementList()
        self.getTextStringList()
        super(docxParser, self).__init__(*args, **kwargs)

    def iter_block_items(self, parent):
        if isinstance(parent, _Document):
            parent_elm = parent.element.body
        elif isinstance(parent, _Cell):
            parent_elm = parent._tc
        else:
            raise ValueError("something's not right")

        for child in parent_elm.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def cellCounts(self, t):
        return len(t.columns[0].cells)

    def getTextStringList(self):
        for block in self.iter_block_items(self.document):
            if block.style.name == 'Heading 3 not-in-TOC':
                if block.text not in self.textStringList:
                    self.textStringList.append(block.text)
        return self.textStringList.sort()
    
    def getRequirementTable(self):
        for i in range(len(self.tables)):
            if self.tables[i].rows[0].cells[0].text == 'SRS Requirement':
                self.sysRqmtTable = self.tables[i]
            elif self.tables[i].rows[0].cells[0].text == 'Requirement*':
                self.srsRqmtTable.append(self.tables[i])
            elif self.tables[i].rows[0].cells[0].text == 'Action/Graphic':
                self.srsRqmtTable1.append(self.tables[i])

    def getSrsRequirementList(self):
        for p in self.document.paragraphs:
            if 'Rqmt-2-Name' in p.style.name or 'Rqmt-3-Name' in p.style.name or 'Rqmt-4-Name' in p.style.name:
                if p.text not in self.srsRequirementList:
                    self.srsRequirementList.append(p.text)
        for i in range(len(self.srsRqmtTable)):
            cell_counts = self.cellCounts(self.srsRqmtTable[i])
            column = self.srsRqmtTable[i].columns[0].cells
            for j in range(cell_counts):
                if 'Rqmt-0-Name' in column[j].paragraphs[0].style.name:
                    if column[j].text not in self.srsRequirementList:
                        self.srsRequirementList.append(column[j].text)
        for i in range(len(self.srsRqmtTable1)):
            cell_counts = self.cellCounts(self.srsRqmtTable1[i])
            column = self.srsRqmtTable1[i].columns[0].cells
            for j in range(cell_counts):
                paragraph_counts = len(column[j].paragraphs)
                for k in range(paragraph_counts):
                    if 'Rqmt-0-Name' in column[j].paragraphs[k].style.name:
                        if column[j].paragraphs[k].text not in self.srsRequirementList:
                            self.srsRequirementList.append(column[j].paragraphs[k].text)
        return self.srsRequirementList.sort()
    
    def getSystemRequirementList(self):
        cell_counts = self.cellCounts(self.sysRqmtTable)
        column1 = self.sysRqmtTable.columns[1].cells
        column2 = self.sysRqmtTable.columns[2].cells
        for i in range(cell_counts-1):
            if column1[i+1].text not in self.systemRequirementList and column1[i+1].text != 'na':
                s = column1[i+1].text.split('\n')
                for element in s:
                    if element and element not in self.systemRequirementList: 
                        self.systemRequirementList.append(element)
        for i in range(cell_counts-1):
            if column2[i+1].text not in self.systemRequirementList and column2[i+1].text != 'na':
                s = column2[i+1].text.split('\n')
                for element in s:
                    if element and element not in self.systemRequirementList: 
                        self.systemRequirementList.append(element)
        return self.systemRequirementList.sort()
   
    def printList(self, list):
        for i in list:
            print i
    def getPageNumber(self):
        f = open('list.txt', 'r')
        lines = f.readlines()
        f.close()
        pageList = list()
        for line in lines:
            if line[0] == '3' and line[3] == '.' or line[4] == '.':
                protocolNameWithPage = line[line.find('\t'):].replace('\t', ' ').strip()
                pageList.append(protocolNameWithPage)



##path = 'C:\\Users\\bing.ma\\Desktop\\Dry_INOmax-3-1-6\\Support\\31122-rev-16-INOmax-DSIR-Plus-MRI-SRS.docx'
##path = os.path.join(raw_input("Type or copy/paste the SRS file folder path.\r\n"))    
#path = r'C:\Users\bing.ma\Desktop\go-into-report-blinding\32817-FinalReport-Blinding-Draft.docx'






