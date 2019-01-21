from docx import Document
from docx.opc import exceptions

def DocxWrite(title, tableContents):
    try:
        document = Document()
    except exceptions.PackageNotFoundError:
        import os
        import sys
        dir_path = os.getcwd()
        full_path = os.path.join(dir_path, 'docx\templates\default.docx')
        print full_path
        document = Document(full_path)
        
    document.add_heading(title, 0)
    table = document.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Requirement'
    hdr_cells[1].text = 'Test Protocol Results Names'
    #table = document.add_table()
    try:
        with open("page_number.txt", 'r') as f:
            page_number_file = "page_number.txt"
            addPageNum(tableContents, page_number_file)
    except IOError:
        pass
    for key in sorted(tableContents):
        row_cells = table.add_row().cells
        row_cells[0].text = str(key)
        if type(tableContents[key]) is list:
            row_cells[1].text = ', '.join(tableContents[key])
        else:
            row_cells[1].text = tableContents[key]
    document.save( title + '.docx')

def addPageNum(lst, file):
    f = open(file, 'r')
    lines = f.readlines()
    f.close()
    pageNum = [] 
    for line in lines:
        try:
            if line[0] == '3' and line[3] == '.' or line[4] == '.':
                pageNum.append(line[line.find('\t')+len('\t'):].replace('\t', ' ').rstrip())
        except IndexError:
            pass
    for key in lst:
        if type(lst[key]) is list:
            for i in range(len(lst[key])):
                for j in pageNum:
                    if lst[key][i] in j:
                        lst[key][i] = j
        else:
            for j in pageNum:
                if lst[key] in j:
                    lst[key] = j





        
            


    
